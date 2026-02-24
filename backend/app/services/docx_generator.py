"""
DOCX Generator Service — Generates Microsoft Word documents locally.
"""

import os
import logging
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

class DocxGeneratorService:
    def __init__(self, output_dir):
        self.output_dir = output_dir
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)

    def generate(self, project, filename=None):
        """Generate a .docx file from project data."""
        if not filename:
            filename = f"{project.get('_id', 'story')}.docx"
            
        filepath = os.path.join(self.output_dir, filename)
        
        try:
            doc = Document()
            
            # Title
            title = doc.add_heading(project.get("title", "My Story"), 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Subtitle
            if project.get("subtitle"):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(project.get("subtitle"))
                run.italic = True
                run.font.size = Pt(14)
            
            doc.add_page_break()
            
            # Chapters
            for chapter in project.get("chapters", []):
                doc.add_heading(chapter.get("chapter_title", "Untitled Chapter"), level=1)
                
                # Content
                paragraphs = chapter.get("content", "").split("\n")
                for para in paragraphs:
                    if para.strip():
                        doc.add_paragraph(para.strip())
                
                doc.add_page_break()
            
            doc.save(filepath)
            return filepath
            
        except Exception as e:
            logger.error(f"DOCX generation error: {e}")
            raise Exception(f"Failed to generate DOCX: {str(e)}")
